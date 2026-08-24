from io import BytesIO
from typing import Any, Dict, Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


GREEN = "234E32"
LIGHT_GREEN = "EAF3E8"
INK = "1F2923"
MUTED = "5E6A62"
LINE = "D8E1D7"
WARNING = "FFF4D6"


def _shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=110, start=130, bottom=110, end=130) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _keep_row_together(row) -> None:
    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    properties.append(cant_split)


def _remove_paragraph_border(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    border = properties.find(qn("w:pBdr"))
    if border is not None:
        properties.remove(border)


def _remove_style_border(style) -> None:
    properties = style.element.get_or_add_pPr()
    border = properties.find(qn("w:pBdr"))
    if border is not None:
        properties.remove(border)


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, before, after in (
        ("Title", 28, 0, 10),
        ("Subtitle", 12, 0, 16),
        ("Heading 1", 18, 20, 8),
        ("Heading 2", 13, 14, 5),
        ("Heading 3", 11, 10, 4),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Subtitle" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(GREEN if name != "Subtitle" else MUTED)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")
        if name == "Title":
            _remove_style_border(style)


def _configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    header = section.header.paragraphs[0]
    header.text = "VALINOR  |  REGISTRO AUDITÁVEL"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.runs[0]
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Documento gerado a partir do registro imutável do procedimento  |  ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    for run in footer.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def _add_label(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GREEN)


def _add_callout(document: Document, title: str, text: str, warning: bool = False) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.7)
    _keep_row_together(table.rows[0])
    cell = table.cell(0, 0)
    _shade(cell, WARNING if warning else LIGHT_GREEN)
    _set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN)
    paragraph = cell.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(0)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def _add_key_value_table(document: Document, rows: Iterable[tuple[str, str]]) -> None:
    data = list(rows)
    table = document.add_table(rows=len(data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.65)
    table.columns[1].width = Inches(5.05)
    for index, (label, value) in enumerate(data):
        left, right = table.rows[index].cells
        left.width = Inches(1.65)
        right.width = Inches(5.05)
        _shade(left, LIGHT_GREEN)
        for cell in (left, right):
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        left.text = label
        right.text = value or "Não informado"
        left.paragraphs[0].runs[0].bold = True
        left.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(GREEN)
        for paragraph in (left.paragraphs[0], right.paragraphs[0]):
            paragraph.paragraph_format.space_after = Pt(0)


def _format_timestamp(value: Any) -> str:
    if not value:
        return "Não registrado"
    return str(value).replace("T", " ").replace("+00:00", " UTC")


def _join_items(items: Any) -> str:
    if not items:
        return "Não registrado."
    if isinstance(items, list):
        return "; ".join(str(item) for item in items)
    return str(items)


def _terms_reference(entry: Any) -> str:
    """Versão e hash do texto que a parte aceitou: é o que permite conferir
    depois exatamente o que foi acordado."""
    entry = entry or {}
    version = entry.get("terms_version")
    digest = entry.get("terms_sha256")
    if not version or not digest:
        return "Não registrado"
    return f"{version} (SHA-256 {digest})"


AI_STAGES = (
    ("Conciliação", "conciliation"),
    ("Organização", "organized"),
    ("Decisão", "decision"),
    ("Auditoria", "review"),
)


def _stage_provenance(case: Dict[str, Any]) -> list[tuple[str, str]]:
    """Com qual modelo e com qual versão de prompt cada etapa foi produzida."""
    rows = []
    for label, field in AI_STAGES:
        stage = case.get(field) or {}
        execution = stage.get("execution") or {}
        if not execution:
            continue
        prompt = execution.get("prompt") or {}
        model = execution.get("model") or "não executado por IA"
        prompt_text = (
            f" | prompt {prompt.get('version')} ({prompt.get('sha256')})"
            if prompt.get("version")
            else ""
        )
        drift = " | PROMPT DIVERGENTE DO MANIFESTO" if execution.get("prompt_drift") else ""
        rows.append((label, f"{model}{prompt_text}{drift}"))
    return rows


def build_docx_report(case: Dict[str, Any]) -> BytesIO:
    document = Document()
    _configure_styles(document)
    _configure_page(document)

    _add_label(document, "Relatório final do procedimento")
    title = document.add_heading(case.get("title") or "Caso sem título", 0)
    _remove_paragraph_border(title)
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run(
        f"Caso {case.get('id', '')}  |  {case.get('claimant', '')} contra {case.get('respondent', '')}"
    )
    _add_callout(
        document,
        "Integridade verificável",
        "Este relatório compila o que ocorreu no procedimento. Documentos, manifestações e eventos são identificados por autoria, data e hash sempre que disponível.",
    )
    _add_key_value_table(
        document,
        (
            ("Status", case.get("status", "")),
            ("Criado em", _format_timestamp(case.get("created_at"))),
            ("Atualizado em", _format_timestamp(case.get("updated_at"))),
            ("Cliente", case.get("claimant", "")),
            ("Empresa", case.get("respondent", "")),
        ),
    )

    document.add_heading("1. Garantias do procedimento", level=1)
    consent = case.get("consent") or {}
    contradictory = case.get("contradictory") or {}
    _add_key_value_table(
        document,
        (
            ("Aceite bilateral", "Completo" if consent.get("complete") else "Pendente"),
            ("Aceite do cliente", _format_timestamp((consent.get("claimant") or {}).get("accepted_at"))),
            ("Termos aceitos pelo cliente", _terms_reference(consent.get("claimant"))),
            ("Aceite da empresa", _format_timestamp((consent.get("respondent") or {}).get("accepted_at"))),
            ("Termos aceitos pela empresa", _terms_reference(consent.get("respondent"))),
            ("Contraditório", "Completo" if contradictory.get("complete") else "Pendente"),
            ("Manifesto", "Travado" if case.get("manifest_locked") else "Ainda aberto"),
        ),
    )

    document.add_heading("2. Participantes, convites e prazos", level=1)
    participants = case.get("participants") or []
    if participants:
        for participant in participants:
            document.add_paragraph(
                f"{participant.get('display_name')} ({participant.get('email')}) - papel: {participant.get('role')}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("O caso utiliza credenciais locais por papel; nenhuma conta foi vinculada.")
    deadlines = case.get("deadlines") or []
    if deadlines:
        document.add_heading("Agenda processual", level=2)
        for deadline in deadlines:
            document.add_paragraph(
                f"{deadline.get('label')} - {deadline.get('assigned_to')} - {_format_timestamp(deadline.get('due_at'))} - {deadline.get('status')}",
                style="List Bullet",
            )

    document.add_heading("3. Materiais e contraditório", level=1)
    documents = case.get("documents") or []
    if not documents:
        document.add_paragraph("Nenhum material foi juntado ao procedimento.")
    for index, item in enumerate(documents, start=1):
        document.add_heading(f"3.{index} {item.get('name', 'Documento')}", level=2)
        _add_key_value_table(
            document,
            (
                ("Autor", item.get("submitted_by", "")),
                ("Tipo", item.get("material_type", "")),
                ("Finalidade", item.get("purpose", "")),
                ("Hash SHA-256", item.get("sha256", "")),
                ("Disponibilizado", _format_timestamp(item.get("disclosed_at"))),
                ("Ciência", _format_timestamp(item.get("acknowledged_at"))),
                ("Manifestação", item.get("response_status", "pending")),
                ("Admitido", "Sim" if item.get("admitted") else "Não"),
            ),
        )
        if item.get("response_text"):
            document.add_paragraph(f"Resposta da contraparte: {item['response_text']}")

    document.add_heading("4. Tentativas de conciliação ou mediação", level=1)
    rounds = case.get("conciliation_rounds") or []
    if not rounds:
        document.add_paragraph("Nenhuma rodada foi registrada.")
    for round_data in rounds:
        number = round_data.get("round_number", "-")
        document.add_heading(f"Rodada {number}", level=2)
        document.add_paragraph(
            round_data.get("rationale")
            or round_data.get("proposal_rationale")
            or round_data.get("stop_reason")
            or "Rodada registrada sem síntese textual."
        )
        _add_key_value_table(
            document,
            (
                ("Caminho", str(round_data.get("recommended_path", ""))),
                ("Convergência", str(round_data.get("convergence", ""))),
                ("Continuar", "Sim" if round_data.get("continue_recommended") else "Não"),
                ("Próximo foco", str(round_data.get("next_round_focus", ""))),
            ),
        )

    document.add_heading("5. Organização e decisão por IA", level=1)
    organized = case.get("organized") or {}
    decision = case.get("decision") or {}
    review = case.get("review") or {}
    if organized:
        document.add_heading("Síntese organizada", level=2)
        document.add_paragraph(str(organized.get("factual_overview") or organized.get("summary") or organized))
        for label, key in (("Fatos controvertidos", "disputed_facts"), ("Questões", "issues")):
            if organized.get(key):
                document.add_paragraph(f"{label}: {_join_items(organized.get(key))}")
    if decision:
        document.add_heading("Decisão", level=2)
        _add_callout(
            document,
            "Resultado computacional",
            str(decision.get("decision") or decision.get("outcome") or "Sem resultado registrado."),
            warning=bool(decision.get("requires_human_review")),
        )
        if decision.get("reasoning"):
            document.add_paragraph(_join_items(decision["reasoning"]))
    else:
        document.add_paragraph("A decisão por IA ainda não foi gerada.")
    if review:
        document.add_heading("Auditoria independente", level=2)
        _add_key_value_table(
            document,
            (
                ("Aprovada", "Sim" if review.get("approved") else "Não"),
                ("Revisão humana", "Necessária" if review.get("requires_human_review") else "Não indicada"),
                ("Conclusão", str(review.get("summary") or review.get("review") or "")),
            ),
        )

    document.add_heading("6. Manifesto e trilha de auditoria", level=1)
    manifest = case.get("locked_manifest") or {}
    _add_key_value_table(
        document,
        (
            ("Hash do manifesto", str(manifest.get("manifest_hash", "Não disponível"))),
            ("Assinatura", str(manifest.get("platform_signature", "Não disponível"))),
            ("Algoritmo", str(manifest.get("signature_algorithm", "Não disponível"))),
        ),
    )
    provenance = _stage_provenance(case)
    if provenance:
        document.add_heading("Procedência das etapas por IA", level=2)
        _add_key_value_table(document, provenance)
        if any("DIVERGENTE" in value for _, value in provenance):
            _add_callout(
                document,
                "Divergência de prompt",
                "Uma ou mais etapas rodaram com prompt diferente do fixado no "
                "manifesto travado. O resultado exige conferência humana.",
                warning=True,
            )

    events = case.get("audit_log") or []
    document.add_heading(f"Eventos registrados ({len(events)})", level=2)
    for event in events:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(f"{_format_timestamp(event.get('timestamp_utc'))} - ").bold = True
        paragraph.add_run(f"{event.get('event_type')} | hash {event.get('event_hash', '')}")

    document.add_heading("7. Nota de uso", level=1)
    _add_callout(
        document,
        "Limite jurídico",
        "O sistema profere uma decisão computacional conforme o procedimento configurado. A saída não constitui, por si só, sentença arbitral ou decisão estatal; eventual eficácia jurídica depende da estrutura contratual adotada, do consentimento válido e da legislação aplicável.",
        warning=True,
    )

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
